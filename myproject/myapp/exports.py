import csv
import io
from datetime import datetime
from django.http import HttpResponse
from django.template.loader import render_to_string
from playwright.sync_api import sync_playwright
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_csv(queryset, filename, headers, data_func):
    """
    Exports a queryset to CSV.
    data_func: function that takes an object and returns a list of values matching headers.
    """
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
    
    writer = csv.writer(response)
    writer.writerow(headers)
    
    for obj in queryset:
        writer.writerow(data_func(obj))
        
    return response

def export_to_excel(queryset, filename, headers, data_func):
    """
    Exports a queryset to Excel (.xlsx).
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Data Export"
    
    # Styling headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = Alignment(horizontal="center", vertical="center")
    
    # Note: openpyxl Fill is a bit more complex, using basic headers for now
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True)
        # Auto-fit width (approximate)
        ws.column_dimensions[ws.cell(row=1, column=col_num).column_letter].width = len(header) + 5

    for row_num, obj in enumerate(queryset, 2):
        row_data = data_func(obj)
        for col_num, value in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col_num, value=str(value))

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
    wb.save(response)
    return response

def export_to_word(queryset, filename, title, headers, data_func):
    """
    Exports a queryset to Word (.docx).
    """
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make bold
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    for obj in queryset:
        row_cells = table.add_row().cells
        row_data = data_func(obj)
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
    doc.save(response)
    return response

def export_to_pdf(queryset, filename, template_name, title, headers, data_func):
    """
    Exports a queryset to PDF using Playwright.
    """
    # Prepare data for template
    data_rows = []
    for obj in queryset:
        data_rows.append(data_func(obj))
        
    context = {
        'title': title,
        'headers': headers,
        'rows': data_rows,
        'date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    }
    
    html_string = render_to_string(template_name, context)
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html_string)
        # Use A4, landscape might be better for many columns
        pdf_bytes = page.pdf(format="A4", print_background=True, margin={"top": "1cm", "bottom": "1cm", "left": "1cm", "right": "1cm"})
        browser.close()
    
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
    return response
